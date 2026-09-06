"""Unit and integration tests for TCD Word (.docx) exam export."""

import io
import pathlib
import zipfile

import docx
import docxtpl
import pytest

from quizml.quizmlyaml import loads
from quizml.renderer.docx import render_docx
from quizml.templates.create_tcd_docx_template import create_tcd_docx_template


@pytest.fixture(scope="module")
def tcd_docx_template_path():
    """Returns path to the packaged tcd-exam.docx template."""
    from quizml.filelocator import locate

    return pathlib.Path(locate.path("tcd-exam.docx"))


def test_tcd_docx_template_variables(tcd_docx_template_path):
    """Verifies that tcd-exam.docx has the expected Jinja2 variables."""
    assert tcd_docx_template_path.exists()
    tpl = docxtpl.DocxTemplate(str(tcd_docx_template_path))
    vars = tpl.get_undeclared_template_variables()
    assert "header" in vars
    assert "questions" in vars
    assert "solutions" in vars


def test_tcd_docx_generator_script(tmp_path):
    """Verifies that create_tcd_docx_template can build a valid docx template."""
    repo_root = pathlib.Path(__file__).parent.parent
    source = repo_root / "tmp-frcs" / "exam-paper-template-reassessment.docx"
    if not source.exists():
        pytest.skip("Base reassessment template not found in workspace")

    target = tmp_path / "generated-tcd-exam.docx"
    create_tcd_docx_template(source, target)
    assert target.exists()
    assert target.stat().st_size > 50000

    tpl = docxtpl.DocxTemplate(str(target))
    vars = tpl.get_undeclared_template_variables()
    assert "header" in vars
    assert "questions" in vars


def test_tcd_docx_cover_page_rendering(tcd_docx_template_path):
    """Verifies rendering of all cover page metadata fields."""
    yaml_text = """
title: 'Telecommunications Exam'
modulecode: 'EEU33C05'
modulename: 'Telecommunications'
programmename: 'Electronic Engineering'
programmeyearname: 'Junior Sophister'
examsemester: 'Semester 1'
examyear: '2025/26'
examiner: 'Dr. Aleksandra Kaszubowska, Dr. Haraj Bennouri'
instructions: 'Please answer all questions from Section A.'
materials: 'Non-programmable calculators permitted.'
additionalinformation: 'Laplace tables provided.'
---
- type: essay
  marks: 10
  question: 'Discuss modulation schemes.'
"""
    doc_dict, _ = loads(yaml_text)
    rendered_bytes = render_docx(
        {"header": doc_dict["header"], "questions": doc_dict["questions"]},
        tcd_docx_template_path,
    )
    doc = docx.Document(io.BytesIO(rendered_bytes))

    # Verify cover page text
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Faculty of Science, Technology, Engineering and Mathematics" in all_text
    assert "School of Engineering" in all_text
    assert "Electronic and Electrical Engineering" in all_text
    assert "EEU33C05 - Telecommunications" in all_text
    assert "Dr. Aleksandra Kaszubowska, Dr. Haraj Bennouri" in all_text
    assert "Please answer all questions from Section A." in all_text
    assert "Non-programmable calculators permitted." in all_text
    assert "Laplace tables provided." in all_text

    # Verify Table 0 metadata
    assert len(doc.tables) > 0
    tbl = doc.tables[0]
    assert "Electronic Engineering" in tbl.rows[0].cells[0].text
    assert "Junior Sophister" in tbl.rows[1].cells[0].text
    assert "Semester 1" in tbl.rows[1].cells[3].text
    assert "2025/26" in tbl.rows[1].cells[4].text

    # Verify Running Header and Footer
    sec = doc.sections[0]
    header_text = "".join(p.text for p in sec.header.paragraphs)
    assert "EEU33C05" in header_text

    footer_text = "".join(p.text for p in sec.footer.paragraphs)
    assert "Trinity College Dublin" in footer_text
    assert "2025/26" in footer_text


def test_tcd_docx_all_question_types(tcd_docx_template_path):
    """Verifies rendering of all QuizML question types in student and solutions modes."""
    yaml_text = """
title: 'Complete Exam'
modulecode: 'TEST101'
---
- type: mc
  marks: 2.0
  question: 'What is 2 + 2?'
  choices:
  - o: '3'
  - x: '4'
  - o: '5'

- type: ma
  marks: 3.0
  question: 'Select prime numbers:'
  choices:
  - x: '2'
  - x: '3'
  - o: '4'

- type: tf
  marks: 1.0
  question: 'The Earth is round.'
  answer: true

- type: num
  marks: 2.0
  question: 'What is pi to 2 decimals?'
  answer: 3.14
  tolerance: 0.01

- type: fill
  marks: 2.0
  question: 'The capital of Ireland is _____.'
  answers:
  - 'Dublin'

- type: mfill
  marks: 3.0
  question: 'Roses are [color1] and violets are [color2].'
  answers:
    color1: ['red', 'Red']
    color2: ['blue', 'Blue']

- type: matching
  marks: 4.0
  question: 'Match the animals:'
  choices:
  - A: 'Dog'
    B: 'Canine'
  - A: 'Cat'
    B: 'Feline'

- type: essay
  marks: 5.0
  question: 'Explain Fourier Transform.'
  answer: 'Fourier transform decomposes a signal into sinusoids.'
"""
    doc_dict, _ = loads(yaml_text)

    # 1. Student mode
    student_bytes = render_docx(
        {
            "header": doc_dict["header"],
            "questions": doc_dict["questions"],
            "solutions": False,
        },
        tcd_docx_template_path,
    )
    doc_std = docx.Document(io.BytesIO(student_bytes))
    std_text = "\n".join(p.text for p in doc_std.paragraphs)

    assert "Question 1.\t[2.0 marks]" in std_text
    assert "What is 2 + 2?" in std_text
    assert "(a)  3" in std_text
    assert "(b)  4" in std_text
    assert "(c)  5" in std_text
    assert "[X]" not in std_text
    assert "Model Answer" not in std_text
    assert "Acceptable answers" not in std_text
    assert "[oOo]" in std_text

    # 2. Solutions mode
    sol_bytes = render_docx(
        {
            "header": doc_dict["header"],
            "questions": doc_dict["questions"],
            "solutions": True,
        },
        tcd_docx_template_path,
    )
    doc_sol = docx.Document(io.BytesIO(sol_bytes))
    sol_text = "\n".join(p.text for p in doc_sol.paragraphs)

    # MC / MA solutions marking
    assert "[X] (b)  4" in sol_text
    assert "[  ] (a)  3" in sol_text
    assert "[X] (a)  2" in sol_text
    assert "[X] (b)  3" in sol_text
    assert "[  ] (c)  4" in sol_text

    # TF solutions marking
    assert "[X] True" in sol_text

    # Numerical solutions
    assert "Answer: 3.14 (± 0.01)" in sol_text

    # Fill & MFill solutions
    assert "Acceptable answers: Dublin" in sol_text
    assert "• [color1]: red, Red" in sol_text
    assert "• [color2]: blue, Blue" in sol_text

    # Matching pairs
    assert "(a)  Dog\t(1)  Canine" in sol_text
    assert "(b)  Cat\t(2)  Feline" in sol_text

    # Essay model answer
    assert "Model Answer:" in sol_text
    assert "Fourier transform decomposes a signal into sinusoids." in sol_text
    assert "[oOo]" in sol_text


def test_tcd_docx_cli_export(tmp_path):
    """Verifies that quizml CLI builds both docx and docx-solutions targets."""
    import sys

    from quizml.cli.cli import main

    yaml_file = tmp_path / "telecom.yaml"
    yaml_file.write_text(
        """
title: 'Telecoms Quiz'
modulecode: 'EEU33C05'
---
- type: mc
  marks: 2.0
  question: 'What is 10 + 10?'
  choices:
  - o: '15'
  - x: '20'
"""
    )

    sys_argv = sys.argv
    try:
        sys.argv = ["quizml", str(yaml_file), "-t", "docx", "-t", "docx-solutions"]
        main()
    finally:
        sys.argv = sys_argv

    out_docx = tmp_path / "telecom.docx"
    out_sol = tmp_path / "telecom.solutions.docx"

    assert out_docx.exists()
    assert out_sol.exists()

    doc_std = docx.Document(str(out_docx))
    std_text = "\n".join(p.text for p in doc_std.paragraphs)
    assert "(a)  15" in std_text
    assert "[X]" not in std_text

    doc_sol = docx.Document(str(out_sol))
    sol_text = "\n".join(p.text for p in doc_sol.paragraphs)
    assert "[X] (b)  20" in sol_text


def test_tcd_docx_xml_escaping_and_html_sanitization(tcd_docx_template_path):
    """Verifies that HTML tags (<style>, <div>) and unescaped XML characters (<, >, &) do not corrupt the docx."""
    yaml_text = """
title: 'ML Exam & Security'
modulecode: 'CS7001'
---
- type: mc
  marks: 2.0
  question: |
    Here is a confusion matrix:
    <div class="confmat"><style scoped>.confmat { color: red; }</style><table><tr><td>TP</td></tr></table></div>
    What is the condition?
  choices:
  - o: 'FPR < 1% & TPR > 99%'
  - x: 'FPR >= 5% && TPR < 80%'
"""
    doc_dict, _ = loads(yaml_text)
    rendered_bytes = render_docx(
        {
            "header": doc_dict["header"],
            "questions": doc_dict["questions"],
            "solutions": True,
        },
        tcd_docx_template_path,
    )
    # python-docx parsing would raise an XMLSyntaxError if XML nodes are invalid or broken
    doc = docx.Document(io.BytesIO(rendered_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)

    assert "color: red" not in all_text
    assert "<style" not in all_text
    assert "FPR < 1% & TPR > 99%" in all_text
    assert "FPR >= 5% && TPR < 80%" in all_text


def test_tcd_docx_editable_docsecurity(tcd_docx_template_path, tmp_path):
    """Verifies that generated Word documents have DocSecurity=0 so they open as editable, not read-only."""
    import re
    import zipfile

    # 1. Template itself must have DocSecurity 0
    with zipfile.ZipFile(str(tcd_docx_template_path)) as z:
        app_xml = z.read("docProps/app.xml").decode("utf-8")
        assert "<DocSecurity>0</DocSecurity>" in app_xml
        assert "<DocSecurity>4</DocSecurity>" not in app_xml

    # 2. Rendered document must retain DocSecurity 0
    rendered_bytes = render_docx(
        {"header": {"title": "Editable Test"}, "questions": []},
        tcd_docx_template_path,
    )
    with zipfile.ZipFile(io.BytesIO(rendered_bytes)) as z:
        app_xml = z.read("docProps/app.xml").decode("utf-8")
        assert "<DocSecurity>0</DocSecurity>" in app_xml
        assert "<DocSecurity>4</DocSecurity>" not in app_xml
        assert not any("custom" in name.lower() for name in z.namelist())

    # 3. Even if given a template artificially locked with DocSecurity 4, render_docx clears it to 0
    locked_template_path = tmp_path / "locked.docx"
    with zipfile.ZipFile(str(tcd_docx_template_path), "r") as zin, zipfile.ZipFile(
        str(locked_template_path), "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            content = zin.read(item.filename)
            if item.filename == "docProps/app.xml":
                content = content.replace(b"<DocSecurity>0</DocSecurity>", b"<DocSecurity>4</DocSecurity>")
            zout.writestr(item, content)

    # Verify our mock template actually has DocSecurity 4
    with zipfile.ZipFile(str(locked_template_path)) as z:
        assert "<DocSecurity>4</DocSecurity>" in z.read("docProps/app.xml").decode("utf-8")

    # Render through render_docx
    unlocked_bytes = render_docx(
        {"header": {"title": "Unlocked Test"}, "questions": []},
        locked_template_path,
    )
    with zipfile.ZipFile(io.BytesIO(unlocked_bytes)) as z:
        unlocked_xml = z.read("docProps/app.xml").decode("utf-8")
        assert "<DocSecurity>0</DocSecurity>" in unlocked_xml
        assert "<DocSecurity>4</DocSecurity>" not in unlocked_xml


def test_tcd_docx_multiline_instructions_and_headers(tcd_docx_template_path):
    """Verifies that multiline header metadata and instructions do not produce nested <w:p> tags."""
    import xml.etree.ElementTree as ET
    from quizml.transcoder import MarkdownTranscoder

    yaml_text = """
title: 'EEU44C16 Final Exam'
modulecode: 'EEU44C16'
modulename: |
  Deep Learning and its Applications

  Final Exam
examiner: 'Dr François Pitié'
examyear: 2024
examdate: 09-Dec-24
examtime: 09:30-11:30
examvenue: Campus
instructions: |
  Answer all questions.

  For the MCQ questions, there is no negative marks.

  Return both the answer sheet and your scripts.
materials: |
  Calculators
---
- type: mc
  marks: 2.5
  question: 'What is deep learning?'
  choices:
  - x: 'A subset of ML'
  - o: 'A database'
"""
    doc_dict, _ = loads(yaml_text)
    transcoder = MarkdownTranscoder(doc_dict)
    target = {"name": "docx", "template": str(tcd_docx_template_path), "fmt": "docx"}
    yaml_transcoded = transcoder.transcode_target(target)

    rendered_bytes = render_docx(yaml_transcoded, tcd_docx_template_path)

    # 1. OpenXML structure must have 0 nested <w:p> elements
    with zipfile.ZipFile(io.BytesIO(rendered_bytes)) as z:
        doc_xml = z.read("word/document.xml")
        root = ET.fromstring(doc_xml)
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        nested_p = []
        for p in root.iter(f"{{{W_NS}}}p"):
            for child in p.iter(f"{{{W_NS}}}p"):
                if child is not p:
                    nested_p.append((p, child))
        assert len(nested_p) == 0, f"Found {len(nested_p)} nested <w:p> elements"

    # 2. python-docx can open and read separate instruction paragraphs
    doc = docx.Document(io.BytesIO(rendered_bytes))
    all_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "Answer all questions." in all_paras
    assert "For the MCQ questions, there is no negative marks." in all_paras
    assert "Return both the answer sheet and your scripts." in all_paras
    assert "Calculators" in all_paras


def test_tcd_docx_answer_sheet_student_and_solutions(tcd_docx_template_path):
    """Verifies answer sheet generation, candidate identification box, and bubbles."""
    yaml_text = """
title: 'Answer Sheet Verification Exam'
modulecode: 'EEU44C16'
---
- type: mc
  marks: 2.5
  question: 'Question 1'
  choices:
  - o: 'Choice 1'
  - x: 'Choice 2'
  - o: 'Choice 3'

- type: tf
  marks: 1.0
  question: 'Question 2'
  answer: true

- type: essay
  marks: 5.0
  question: 'Question 3'
"""
    doc_dict, _ = loads(yaml_text)

    # 1. Student Mode
    std_bytes = render_docx(
        {"header": doc_dict["header"], "questions": doc_dict["questions"], "solutions": False},
        tcd_docx_template_path,
    )
    doc_std = docx.Document(io.BytesIO(std_bytes))

    # Check candidate table
    cand_tbl = doc_std.tables[1]
    assert "Exam Number" in cand_tbl.rows[0].cells[0].text
    assert "Seat Number" in cand_tbl.rows[0].cells[0].text
    assert "Mark your exam number below" in cand_tbl.rows[0].cells[1].text
    assert "⓪ ① ② ③ ④ ⑤ ⑥ ⑦ ⑧ ⑨" in cand_tbl.rows[0].cells[1].text

    # Check answer banner
    all_text = "\n".join(p.text for p in doc_std.paragraphs)
    assert "All your MCQ answers must be filled in on this answer page." in all_text
    assert "This page was intentionally left blank." in all_text

    # Check grid table in student mode
    grid_tbl = doc_std.tables[2]
    assert len(grid_tbl.rows) == 1  # 3 questions -> 1 row of 3 columns
    row0 = [c.text for c in grid_tbl.rows[0].cells]
    assert "Q.1\tⒶ Ⓑ Ⓒ" in row0[0]
    assert "Q.2\tⓉ Ⓕ" in row0[1]
    assert "Q.3\tessay question" in row0[2]

    # 2. Solutions Mode
    sol_bytes = render_docx(
        {"header": doc_dict["header"], "questions": doc_dict["questions"], "solutions": True},
        tcd_docx_template_path,
    )
    doc_sol = docx.Document(io.BytesIO(sol_bytes))
    grid_tbl_sol = doc_sol.tables[2]
    row0_sol = [c.text for c in grid_tbl_sol.rows[0].cells]
    assert "Q.1\tⒶ 🅑 Ⓒ" in row0_sol  # Choice 2 is marked correct
    assert "Q.2\t🅣 Ⓕ" in row0_sol  # True is marked correct


def test_tcd_docx_answer_sheet_midterm_mode(tcd_docx_template_path):
    """Verifies candidate box switches to Student Name/Number and 8 digit rows on midterm."""
    yaml_text = """
title: 'Midterm Exam'
modulecode: 'EEU44C16'
midterm: true
---
- type: mc
  marks: 2.0
  question: 'Midterm Q1'
  choices:
  - x: 'Right'
  - o: 'Wrong'
"""
    doc_dict, _ = loads(yaml_text)
    rendered_bytes = render_docx(
        {"header": doc_dict["header"], "questions": doc_dict["questions"]},
        tcd_docx_template_path,
    )
    doc = docx.Document(io.BytesIO(rendered_bytes))
    cand_tbl = doc.tables[1]
    left_cell = cand_tbl.rows[0].cells[0].text
    right_cell = cand_tbl.rows[0].cells[1].text

    assert "Student Name" in left_cell
    assert "Student Number" in left_cell
    assert "Mark your student number below" in right_cell
    # 8 rows of digits
    assert right_cell.count("⓪ ① ② ③ ④ ⑤ ⑥ ⑦ ⑧ ⑨") == 8




