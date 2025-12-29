from docx import Document
from docxtpl import DocxTemplate
import sys

def create_template(filename="template.docx"):
    doc = Document()
    
    doc.add_heading('{{ header.title }}', 0)
    
    doc.add_paragraph('Duration: {{ header.duration }} minutes')
    
    doc.add_paragraph('Instructions: {{ header.instructions }}')
    
    doc.add_heading('Questions', level=1)
    
    # Add a loop over questions
    # Note: In docxtpl, loops are often done with {% for ... %} inside the document text
    # or by using table rows. Here we'll just add text with jinja tags.
    
    p = doc.add_paragraph('{% for q in questions %}')
    
    doc.add_heading('Question {{ loop.index }}', level=2)
    
    # We assume 'question' might contain some markdown/html, but for now just dump it
    # docxtpl can handle RichText if we passed it, but here we just put the variable
    doc.add_paragraph('{{ q.question }}')
    
    # Choices
    p = doc.add_paragraph('{% if q.choices %}')
    p = doc.add_paragraph('{% for c in q.choices %}')
    # Handling the complex choice structure (list of dicts)
    # This is a simplification, might need adjustment based on exact data structure
    p = doc.add_paragraph(' - {{ c }}') 
    p = doc.add_paragraph('{% endfor %}')
    p = doc.add_paragraph('{% endif %}')
    
    p = doc.add_paragraph('{% endfor %}')
    
    doc.save(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    try:
        import docxtpl
    except ImportError:
        print("Please install docxtpl: pip install docxtpl")
        sys.exit(1)
        
    create_template("quiz_template.docx")
