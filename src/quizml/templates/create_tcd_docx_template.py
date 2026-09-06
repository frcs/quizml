"""Generates the official TCD Word exam template (tcd-exam.docx) from the institutional template.

Preserves the official Trinity College Dublin crest/logo, margins, typography (Calibri 12pt),
running headers/footers with dynamic page numbers and copyright, cover page metadata layout,
and inserts Jinja2 tags supporting all QuizML question types for student and solution modes.
"""

from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _unwrap_sdt_elements(parent_element):
    """Unwraps all <w:sdt> structured document tags into their inner content elements."""
    for el in list(parent_element):
        tag = el.tag.split("}")[-1]
        if tag == "sdt":
            sdt_content = el.find(f"{{{W_NS}}}sdtContent")
            if sdt_content is not None:
                children = list(sdt_content)
                idx = parent_element.index(el)
                for child in children:
                    parent_element.insert(idx, child)
                    idx += 1
                parent_element.remove(el)
        elif len(el) > 0:
            _unwrap_sdt_elements(el)


def _set_table_borders_none(table):
    """Removes all visible borders from an OpenXML table."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tblPr = table._tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for b_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{b_name}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)


def create_tcd_docx_template(
    source_docx_path: str | Path, output_docx_path: str | Path
) -> Path:
    """Transforms the TCD reassessment Word template into a QuizML Jinja2 docxtpl template."""
    source_path = Path(source_docx_path)
    output_path = Path(output_docx_path)

    if not source_path.exists():
        raise FileNotFoundError(f"Source template not found: {source_path}")

    doc = docx.Document(str(source_path))

    # 1. Unwrap all <w:sdt> structured document tags in body and tables
    _unwrap_sdt_elements(doc.element.body)

    # 2. Configure Cover Page metadata
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt == "Science, Technology, Engineering and Mathematics":
            p.text = (
                "{{ header.faculty or 'Faculty of Science, Technology, Engineering and Mathematics' }}"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.size = Pt(12)
        elif txt == "School of Engineering":
            p.text = "{{ header.school or 'School of Engineering' }}"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.size = Pt(12)
        elif txt == "Electronic and Electrical Engineering":
            p.text = "{{ header.department or 'Electronic and Electrical Engineering' }}"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.name = "Calibri"
                p.runs[0].font.size = Pt(12)
        elif "3C5 - Telecommunications" in txt:
            p.text = (
                "{% if header.modulecode and (header.modulename or header.title) %}"
                "{{ header.modulecode }} - {{ header.modulename or header.title }}"
                "{% else %}{{ header.modulecode or header.modulename or header.title or '' }}{% endif %}"
            )
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(12)
        elif p.style.name == "Heading 5" and ("\t" in p.text or not txt):
            # Date, Venue, Time row
            p.text = (
                "{% if header.examdate or header.examvenue or header.examtime %}"
                "{{ header.examdate or '' }}\t{{ header.examvenue or '' }}\t{{ header.examtime or '' }}"
                "{% endif %}"
            )
        elif txt == "Dr. Aleksandra Kaszubowska, Dr. Haraj Bennouri":
            p.text = "{{ header.examiner }}"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(12)
        elif txt == "Please answer all questions.":
            p.text = "{{p header.instructions }}"
        elif txt.startswith("Non-programmable calculators are permitted"):
            p.text = (
                "{% if header.materials %}{{ header.materials }}"
                "{% else %}Non-programmable calculators are permitted for this examination. "
                "Please indicate the make and model of your calculator on each answer book used.{% endif %}"
            )
        elif txt == "Log Tables permitted":
            p.text = "{{p header.additionalinformation }}"

    # 3. Configure Table 0 (Degree Programme and Session Info)
    if doc.tables:
        tbl = doc.tables[0]
        # Row 0: Programme Name
        if len(tbl.rows) > 0 and len(tbl.rows[0].cells) > 0:
            p_prog = tbl.rows[0].cells[0].paragraphs[0]
            p_prog.text = "{{ header.programmename or header.degree or 'Engineering' }}"
            if p_prog.runs:
                p_prog.runs[0].font.name = "Calibri"
                p_prog.runs[0].font.size = Pt(12)

        # Row 1: Year Level (cell 0), Semester (cell 3), Academic Year (cell 4)
        if len(tbl.rows) > 1:
            row1 = tbl.rows[1]
            if len(row1.cells) > 0 and row1.cells[0].paragraphs:
                p_yr = row1.cells[0].paragraphs[0]
                p_yr.text = (
                    "{{ header.programmeyearname or header.programmeyear or 'Junior Sophister' }}"
                )
                if p_yr.runs:
                    p_yr.runs[0].font.name = "Calibri"
                    p_yr.runs[0].font.size = Pt(12)

            if len(row1.cells) > 3 and row1.cells[3].paragraphs:
                p_sem = row1.cells[3].paragraphs[0]
                p_sem.text = "{{ header.examsemester or 'Semester 1' }}"
                if p_sem.runs:
                    p_sem.runs[0].font.name = "Calibri"
                    p_sem.runs[0].font.size = Pt(12)

            if len(row1.cells) > 4 and row1.cells[4].paragraphs:
                p_ay = row1.cells[4].paragraphs[0]
                p_ay.text = "{{ header.examyear or '2025/26' }}"
                if p_ay.runs:
                    p_ay.runs[0].font.name = "Calibri"
                    p_ay.runs[0].font.size = Pt(12)

    # 4. Configure Running Header and Footer
    for sec in doc.sections:
        if sec.header and sec.header.paragraphs:
            sec.header.paragraphs[0].text = "{{ header.modulecode }}"
            sec.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if sec.header.paragraphs[0].runs:
                sec.header.paragraphs[0].runs[0].bold = True

        if sec.footer and sec.footer.paragraphs:
            p_ftr = sec.footer.paragraphs[0]
            if len(p_ftr.runs) >= 5 and "Trinity" in p_ftr.runs[0].text:
                p_ftr.runs[2].text = " {{ header.examyear or '2026' }}"
                p_ftr.runs[3].text = ""
                p_ftr.runs[4].text = ""
            else:
                for run in p_ftr.runs:
                    if "2026" in run.text:
                        run.text = run.text.replace("2026", "{{ header.examyear or '2026' }}")

    # 5. Locate the Page Break paragraph after cover page
    split_idx = None
    for i, p in enumerate(doc.paragraphs):
        xml = p._element.xml
        if 'w:br w:type="page"' in xml or "pageBreakBefore" in xml:
            split_idx = i
            break

    if split_idx is None:
        raise ValueError("Could not find page break after cover page in source template.")

    # Remove all instructional placeholder paragraphs after the page break
    total_paragraphs = len(doc.paragraphs)
    for _ in range(split_idx + 1, total_paragraphs):
        p = doc.paragraphs[split_idx + 1]
        p._element.getparent().remove(p._element)

    # 6. Insert Blank Page (Page 2 - back of cover page in duplex print)
    p_blank1 = doc.add_paragraph()
    p_blank1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_blank1.paragraph_format.space_before = Pt(250)
    r_b1 = p_blank1.add_run("This page was intentionally left blank.")
    r_b1.italic = True
    r_b1.font.name = "Calibri"
    r_b1.font.size = Pt(12)

    # Page break to start Answer Sheet on Page 3
    p_br1 = doc.add_paragraph()
    p_br1.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

    # 7. Build Answer Sheet (Page 3 - front of sheet 2)
    # Candidate Identification Table (2 columns: ID fields on left, digit bubble matrix on right)
    cand_tbl = doc.add_table(rows=1, cols=2)
    _set_table_borders_none(cand_tbl)
    col_widths_cand = [Inches(3.6), Inches(2.78)]
    for row in cand_tbl.rows:
        for i, w in enumerate(col_widths_cand):
            row.cells[i].width = w

    # Left cell: Exam Number / Seat Number (or Student Name / Number)
    c0 = cand_tbl.rows[0].cells[0]
    p0_1 = c0.paragraphs[0]
    p0_1.paragraph_format.space_before = Pt(4)
    p0_1.paragraph_format.space_after = Pt(14)
    r0_1_label = p0_1.add_run("{{ candidate_info.line1_label }}:\t")
    r0_1_label.font.name = "Calibri"
    r0_1_label.font.size = Pt(11)
    r0_1_line = p0_1.add_run("________________________________")
    r0_1_line.font.name = "Calibri"
    r0_1_line.font.size = Pt(11)

    p0_2 = c0.add_paragraph()
    p0_2.paragraph_format.space_before = Pt(0)
    p0_2.paragraph_format.space_after = Pt(4)
    r0_2_label = p0_2.add_run("{{ candidate_info.line2_label }}:\t")
    r0_2_label.font.name = "Calibri"
    r0_2_label.font.size = Pt(11)
    r0_2_line = p0_2.add_run("________________________________")
    r0_2_line.font.name = "Calibri"
    r0_2_line.font.size = Pt(11)

    # Right cell: Digit bubble grid (0-9)
    c1 = cand_tbl.rows[0].cells[1]
    p1_1 = c1.paragraphs[0]
    p1_1.paragraph_format.space_before = Pt(4)
    p1_1.paragraph_format.space_after = Pt(4)
    r1_1 = p1_1.add_run("{{ candidate_info.mark_instruction }}")
    r1_1.font.name = "Calibri"
    r1_1.font.size = Pt(11)

    p1_2 = c1.add_paragraph()
    p1_2.paragraph_format.space_before = Pt(0)
    p1_2.paragraph_format.space_after = Pt(4)
    p1_2.paragraph_format.line_spacing = 1.15
    r1_2 = p1_2.add_run("{{ candidate_info.digit_rows }}")
    r1_2.font.name = "Calibri"
    r1_2.font.size = Pt(10)

    # Answer Sheet Banner
    p_banner1 = doc.add_paragraph()
    p_banner1.paragraph_format.space_before = Pt(14)
    p_banner1.paragraph_format.space_after = Pt(4)
    r_bn1 = p_banner1.add_run(
        "All your MCQ answers must be filled in on this answer page."
    )
    r_bn1.bold = True
    r_bn1.font.name = "Calibri"
    r_bn1.font.size = Pt(11)

    p_banner2 = doc.add_paragraph()
    p_banner2.paragraph_format.space_before = Pt(0)
    p_banner2.paragraph_format.space_after = Pt(12)
    r_b2_1 = p_banner2.add_run("For ")
    r_b2_1.font.name = "Calibri"
    r_b2_1.font.size = Pt(10.5)

    r_b2_2 = p_banner2.add_run("True")
    r_b2_2.bold = True
    r_b2_2.font.name = "Calibri"
    r_b2_2.font.size = Pt(10.5)

    r_b2_3 = p_banner2.add_run(" or ")
    r_b2_3.font.name = "Calibri"
    r_b2_3.font.size = Pt(10.5)

    r_b2_4 = p_banner2.add_run("False")
    r_b2_4.bold = True
    r_b2_4.font.name = "Calibri"
    r_b2_4.font.size = Pt(10.5)

    r_b2_5 = p_banner2.add_run(
        " questions, mark Ⓣ or Ⓕ. For questions with multiple choices, mark all"
        " solutions that are correct (for instance 🅐🅑Ⓒ🅓)."
    )
    r_b2_5.font.name = "Calibri"
    r_b2_5.font.size = Pt(10.5)

    # 3-Column Answer Grid Table
    grid_tbl = doc.add_table(rows=3, cols=3)
    _set_table_borders_none(grid_tbl)
    col_widths_grid = [Inches(2.12), Inches(2.12), Inches(2.14)]
    for row in grid_tbl.rows:
        for i, w in enumerate(col_widths_grid):
            row.cells[i].width = w

    grid_tbl.rows[0].cells[0].paragraphs[0].text = (
        "{%tr for row in answer_sheet_rows %}"
    )

    for i, col_key in enumerate(["c1", "c2", "c3"]):
        p_col = grid_tbl.rows[1].cells[i].paragraphs[0]
        p_col.paragraph_format.space_before = Pt(1)
        p_col.paragraph_format.space_after = Pt(2)
        p_col.paragraph_format.line_spacing = 1.15
        p_col.paragraph_format.tab_stops.add_tab_stop(Inches(0.40))
        r_num = p_col.add_run(f"{{{{ row.{col_key}.num }}}}")
        r_num.bold = True
        r_num.font.name = "Calibri"
        r_num.font.size = Pt(10.5)
        r_sep = p_col.add_run(f"{{{{ row.{col_key}.sep }}}}")
        r_sep.font.name = "Calibri"
        r_sep.font.size = Pt(10.5)
        r_bub = p_col.add_run(f"{{{{ row.{col_key}.bubbles }}}}")
        r_bub.font.name = "Calibri"
        r_bub.font.size = Pt(10.5)

    grid_tbl.rows[2].cells[0].paragraphs[0].text = "{%tr endfor %}"

    # Page break after Answer Sheet
    p_br2 = doc.add_paragraph()
    p_br2.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

    # 8. Insert Blank Page (Page 4 - back of detachable answer sheet in duplex print)
    p_blank2 = doc.add_paragraph()
    p_blank2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_blank2.paragraph_format.space_before = Pt(250)
    r_b2 = p_blank2.add_run("This page was intentionally left blank.")
    r_b2.italic = True
    r_b2.font.name = "Calibri"
    r_b2.font.size = Pt(12)

    # Page break to start Questions on Page 5
    p_br3 = doc.add_paragraph()
    p_br3.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

    # 9. Build Questions Loop
    doc.add_paragraph("{%p for q in questions %}")

    # Question Title and Marks Header
    p_q_header = doc.add_paragraph()
    p_q_header.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT)
    r_q = p_q_header.add_run("Question {{ loop.index }}.")
    r_q.bold = True
    r_q.font.name = "Calibri"
    r_q.font.size = Pt(12)

    r_m = p_q_header.add_run("\t[{{ q.marks }} marks]")
    r_m.bold = True
    r_m.font.name = "Calibri"
    r_m.font.size = Pt(12)

    # Question Statement
    doc.add_paragraph("{{p q.question }}")

    # Choice handling for MC and MA
    doc.add_paragraph("{%p if q.choices and q.type in ['mc', 'ma'] %}")
    doc.add_paragraph("{%p for c in q.choices %}")

    # Choice line in solutions mode when correct
    doc.add_paragraph("{%p if solutions and c.x is defined %}")
    p_sol_cor = doc.add_paragraph()
    r_sol_cor = p_sol_cor.add_run(
        "    [X] ({{ 'abcdefghijklmnopqrstuvwxyz'[loop.index0] }})  "
    )
    r_sol_cor.bold = True
    r_sol_cor.font.name = "Calibri"
    r_sol_cor.font.size = Pt(12)
    p_sol_cor.add_run("{{r c.x }}")

    # Choice line in solutions mode when incorrect
    doc.add_paragraph("{%p elif solutions %}")
    p_sol_inc = doc.add_paragraph()
    r_sol_inc = p_sol_inc.add_run(
        "    [  ] ({{ 'abcdefghijklmnopqrstuvwxyz'[loop.index0] }})  "
    )
    r_sol_inc.font.name = "Calibri"
    r_sol_inc.font.size = Pt(12)
    p_sol_inc.add_run("{{r c.o }}")

    # Choice line in regular exam mode
    doc.add_paragraph("{%p else %}")
    p_std = doc.add_paragraph()
    r_std = p_std.add_run(
        "    ({{ 'abcdefghijklmnopqrstuvwxyz'[loop.index0] }})  "
    )
    r_std.font.name = "Calibri"
    r_std.font.size = Pt(12)
    p_std.add_run("{{r c.values() | list | first }}")

    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endif %}")

    # True / False Handling
    doc.add_paragraph("{%p if q.type == 'tf' %}")

    doc.add_paragraph("{%p if solutions and q.answer|string|lower == 'true' %}")
    p_tf_t_sol = doc.add_paragraph()
    r_tf_t1 = p_tf_t_sol.add_run("    [X] True\n    [  ] False")
    r_tf_t1.bold = True
    r_tf_t1.font.name = "Calibri"
    r_tf_t1.font.size = Pt(12)

    doc.add_paragraph("{%p elif solutions %}")
    p_tf_f_sol = doc.add_paragraph()
    r_tf_f1 = p_tf_f_sol.add_run("    [  ] True\n    [X] False")
    r_tf_f1.bold = True
    r_tf_f1.font.name = "Calibri"
    r_tf_f1.font.size = Pt(12)

    doc.add_paragraph("{%p else %}")
    p_tf_std = doc.add_paragraph()
    r_tf_std = p_tf_std.add_run("    (a)  True\n    (b)  False")
    r_tf_std.font.name = "Calibri"
    r_tf_std.font.size = Pt(12)

    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{%p endif %}")

    # Matching Question Handling
    doc.add_paragraph("{%p if q.type == 'matching' and q.choices %}")
    doc.add_paragraph("{%p for c in q.choices %}")
    p_match = doc.add_paragraph()
    p_match.paragraph_format.tab_stops.add_tab_stop(Inches(3.5), WD_TAB_ALIGNMENT.LEFT)
    r_match_a = p_match.add_run(
        "    ({{ 'abcdefghijklmnopqrstuvwxyz'[loop.index0] }})  "
    )
    r_match_a.font.name = "Calibri"
    r_match_a.font.size = Pt(12)
    p_match.add_run("{{r c.A }}")
    r_match_b = p_match.add_run("\t({{ loop.index }})  ")
    r_match_b.font.name = "Calibri"
    r_match_b.font.size = Pt(12)
    p_match.add_run("{{r c.B }}")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endif %}")

    # Numerical Solutions Handling
    doc.add_paragraph("{%p if solutions and q.type == 'num' and q.answer is defined %}")
    p_num = doc.add_paragraph(
        "    Answer: {{ q.answer }}{% if q.tolerance is defined %} (± {{ q.tolerance }}){% endif %}"
    )
    if p_num.runs:
        p_num.runs[0].bold = True
        p_num.runs[0].font.name = "Calibri"
        p_num.runs[0].font.size = Pt(12)
    doc.add_paragraph("{%p endif %}")

    # Fill-in Solutions Handling
    doc.add_paragraph("{%p if solutions and q.type == 'fill' and q.answers %}")
    p_fill = doc.add_paragraph("    Acceptable answers: {{ q.answers | join(', ') }}")
    if p_fill.runs:
        p_fill.runs[0].bold = True
        p_fill.runs[0].font.name = "Calibri"
        p_fill.runs[0].font.size = Pt(12)
    doc.add_paragraph("{%p endif %}")

    # Multi-Fill Cloze Solutions Handling
    doc.add_paragraph("{%p if solutions and q.type == 'mfill' and q.answers %}")
    p_mfill_hdr = doc.add_paragraph("    Acceptable answers:")
    if p_mfill_hdr.runs:
        p_mfill_hdr.runs[0].bold = True
        p_mfill_hdr.runs[0].font.name = "Calibri"
        p_mfill_hdr.runs[0].font.size = Pt(12)
    doc.add_paragraph("{%p for key, values in q.answers.items() %}")
    p_mfill_item = doc.add_paragraph("      • [{{ key }}]: {{ values | join(', ') }}")
    if p_mfill_item.runs:
        p_mfill_item.runs[0].font.name = "Calibri"
        p_mfill_item.runs[0].font.size = Pt(12)
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("{%p endif %}")

    # Essay Model Answer Handling
    doc.add_paragraph("{%p if solutions and q.type == 'essay' and q.answer %}")
    p_essay_hdr = doc.add_paragraph("    Model Answer:")
    if p_essay_hdr.runs:
        p_essay_hdr.runs[0].bold = True
        p_essay_hdr.runs[0].font.name = "Calibri"
        p_essay_hdr.runs[0].font.size = Pt(12)
    doc.add_paragraph("{{p q.answer }}")
    doc.add_paragraph("{%p endif %}")

    # Separation space between questions per guidelines
    doc.add_paragraph()

    # End question loop
    doc.add_paragraph("{%p endfor %}")

    # End of examination marker
    p_end = doc.add_paragraph("[oOo]")
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_end.runs:
        p_end.runs[0].bold = True
        p_end.runs[0].font.name = "Calibri"
        p_end.runs[0].font.size = Pt(12)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    _strip_server_and_security_metadata(output_path)
    print(f"Successfully generated TCD Word template: {output_path}")
    return output_path


def _strip_server_and_security_metadata(docx_path: Path) -> None:
    """Removes SharePoint/server document metadata and ensures DocSecurity is 0 (editable)."""
    import re
    import zipfile

    tmp_file = docx_path.with_suffix(".tmp.docx")
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
            tmp_file, "w", compression=zipfile.ZIP_DEFLATED
        ) as zout:
            for item in zin.infolist():
                name = item.filename
                # Strip all customXml (SharePoint forms, schemas, documentManagement)
                if name.startswith("customXml/"):
                    continue

                # Strip docProps/custom.xml (ContentTypeId linking to SharePoint document library)
                if name == "docProps/custom.xml":
                    continue

                content = zin.read(name)

                if name == "_rels/.rels":
                    text = content.decode("utf-8")
                    text = re.sub(
                        r"<Relationship[^>]*Target=[\"\x27]docProps/custom\.xml[\"\x27][^>]*/>",
                        "",
                        text,
                    )
                    content = text.encode("utf-8")

                elif name == "word/_rels/document.xml.rels":
                    text = content.decode("utf-8")
                    text = re.sub(
                        r"<Relationship[^>]*Target=[\"\x27]\.\./customXml/[^\"\x27]+[\"\x27][^>]*/>",
                        "",
                        text,
                    )
                    content = text.encode("utf-8")

                elif name == "[Content_Types].xml":
                    text = content.decode("utf-8")
                    text = re.sub(
                        r"<Override[^>]*PartName=[\"\x27]/docProps/custom\.xml[\"\x27][^>]*/>",
                        "",
                        text,
                    )
                    text = re.sub(
                        r"<Override[^>]*PartName=[\"\x27]/customXml/[^\"\x27]+[\"\x27][^>]*/>",
                        "",
                        text,
                    )
                    content = text.encode("utf-8")

                elif name == "docProps/app.xml":
                    text = content.decode("utf-8")
                    text = re.sub(
                        r"<DocSecurity>\d+</DocSecurity>",
                        "<DocSecurity>0</DocSecurity>",
                        text,
                    )
                    content = text.encode("utf-8")

                zout.writestr(item, content)

        tmp_file.replace(docx_path)
    finally:
        if tmp_file.exists():
            tmp_file.unlink()


if __name__ == "__main__":
    src_dir = Path(__file__).parent
    repo_root = src_dir.parent.parent.parent
    source = repo_root / "tmp-frcs" / "exam-paper-template-reassessment.docx"
    target = src_dir / "tcd-exam.docx"
    create_tcd_docx_template(source, target)

